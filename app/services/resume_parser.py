"""PDF, Word, and Image text extraction with enhanced quality."""

import os
import re
import logging
import traceback

logger = logging.getLogger(__name__)


class ResumeParserError(Exception):
    """Raised when resume text cannot be extracted."""


def _normalise(text):
    """Clean up OCR / PDF extraction artefacts."""
    # Fix broken words from PDF (e.g. "experi ence" → "experience")
    text = re.sub(r"([a-z])-\n([a-z])", r"\1\2", text)
    # Collapse multiple blank lines to one
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove page-break markers and form-feeds
    text = text.replace("\x0c", "\n")
    # Normalize unicode hyphens / dashes
    text = re.sub(r"[\u2010-\u2015\u2212]", "-", text)
    # Replace smart quotes
    text = re.sub(r"[\u2018\u2019]", "'", text)
    text = re.sub(r"[\u201c\u201d]", '"', text)
    return text.strip()


def extract_text_from_pdf(file_path):
    """Extract plain text from a PDF resume.

    Tries pdfplumber first (better layout), falls back to PyPDF2.
    """
    if not os.path.exists(file_path):
        raise ResumeParserError("File not found.")

    # ── Attempt 1: pdfplumber (best for complex layouts) ─────────────────────
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text(x_tolerance=3, y_tolerance=3)
                if t:
                    texts.append(t)
        if texts:
            return _normalise("\n".join(texts))
    except ImportError:
        pass  # pdfplumber not installed; fall through to PyPDF2
    except Exception:
        pass  # Corrupt or unusual PDF; fall through

    # ── Attempt 2: PyPDF2 ────────────────────────────────────────────────────
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ResumeParserError("PDF is password-protected.") from exc
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        if pages:
            return _normalise("\n".join(pages))
    except ResumeParserError:
        raise
    except ImportError:
        raise ResumeParserError("PyPDF2 is not installed.")
    except Exception as exc:
        raise ResumeParserError(f"Unable to read PDF: {exc}") from exc

    raise ResumeParserError("No readable text found in PDF.")


def extract_text_from_docx(file_path):
    """Extract plain text from a .docx or .doc resume."""
    # ── Attempt 1: python-docx (preserves paragraph structure) ───────────────
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        if paragraphs:
            return _normalise("\n".join(paragraphs))
    except ImportError:
        pass  # python-docx not installed; fall through
    except Exception:
        pass

    # ── Attempt 2: docx2txt (simpler fallback) ────────────────────────────────
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        if text and text.strip():
            return _normalise(text)
    except ImportError:
        pass
    except Exception:
        pass

    raise ResumeParserError("Unable to read Word document. Ensure it is a valid .docx file.")


def extract_text_from_image(file_path):
    """Extract text from an image resume using OCR.
    Tries EasyOCR first, falls back to pytesseract.
    """
    text = ""
    errors = []
    
    # ── Attempt 1: EasyOCR (preferred) ────────────────────────────────────────
    try:
        import easyocr
        # Use English language
        reader = easyocr.Reader(['en'], gpu=False, verbose=False)
        result = reader.readtext(file_path, detail=0)
        text = "\n".join(result)
    except ImportError as e:
        errors.append(f"EasyOCR ImportError: {e}")
        logger.warning(f"EasyOCR not available: {e}")
    except Exception as e:
        errors.append(f"EasyOCR Exception: {e}")
        logger.error(f"EasyOCR extraction failed for {file_path}", exc_info=True)

    if text and text.strip():
        return _normalise(text)

    # ── Attempt 2: pytesseract (fallback) ─────────────────────────────────────
    try:
        from PIL import Image, ImageFilter, ImageEnhance
        import pytesseract

        import os
        if os.path.exists(r'C:\Program Files\Tesseract-OCR\tesseract.exe'):
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        elif os.path.exists(r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'):
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'
        elif os.path.exists(os.path.expandvars(r'%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe')):
            pytesseract.pytesseract.tesseract_cmd = os.path.expandvars(r'%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe')

        img = Image.open(file_path)

        # Convert to RGB if needed (handles RGBA PNGs)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        # Upscale small images to improve OCR accuracy
        w, h = img.size
        if w < 1000:
            scale = 1000 / w
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

        # Convert to greyscale
        img = img.convert("L")

        # Sharpen
        img = img.filter(ImageFilter.SHARPEN)

        # Increase contrast
        img = ImageEnhance.Contrast(img).enhance(2.0)

        # OCR with best accuracy config
        custom_config = r"--oem 3 --psm 6"
        text = pytesseract.image_to_string(img, lang="eng", config=custom_config)

    except ImportError as e:
        errors.append(f"PyTesseract/PIL ImportError: {e}")
        logger.warning(f"PyTesseract/PIL not available: {e}")
    except Exception as e:
        errors.append(f"PyTesseract Exception: {e}")
        logger.error(f"PyTesseract extraction failed for {file_path}", exc_info=True)
        
    if not text or not text.strip():
        error_msg = "Unable to read image resume. OCR failed. Reasons: " + " | ".join(errors)
        logger.error(error_msg)
        raise ResumeParserError(error_msg)

    return _normalise(text)


def extract_text(file_path, file_type):
    """Route extraction by file type."""
    ft = (file_type or "").lower().strip(".")
    if ft == "pdf":
        return extract_text_from_pdf(file_path)
    elif ft in ("doc", "docx"):
        return extract_text_from_docx(file_path)
    elif ft in ("jpg", "jpeg", "png"):
        return extract_text_from_image(file_path)
    else:
        raise ResumeParserError(f"Unsupported file type: {file_type}")
